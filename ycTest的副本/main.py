import os
import re
import sys
import subprocess
from docx import Document

# 自动安装依赖


def remove_bracket_content(text):
    # 只删除括号内内容，保留括号
    return re.sub(r'（[^（）]*）', '（）', text)

# 判断是否为题干段落（以数字+点+空格开头）
def is_question_stem(text):
    # 例如：1. xxx  12. xxx  8. xxx
    return bool(re.match(r'^\d+\.\s*', text.strip()))

def process_docx_file(filepath):
    doc = Document(filepath)
    new_paragraphs = []
    first_stem = True
    for idx, para in enumerate(doc.paragraphs):
        orig_text = para.text
        is_stem = is_question_stem(orig_text)
        if is_stem:
            if not first_stem:
                new_paragraphs.append('')  # 在新题干前插入空行
            first_stem = False
            if '（' in orig_text and '）' in orig_text:
                new_text = remove_bracket_content(orig_text)
            else:
                new_text = orig_text
        else:
            new_text = orig_text
        new_paragraphs.append(new_text)
    # 重新生成文档
    new_doc = Document()
    for text in new_paragraphs:
        new_doc.add_paragraph(text)
    new_filepath = filepath.replace('.docx', '_new.docx')
    new_doc.save(new_filepath)

def main():
    files_dir = os.path.join(os.path.dirname(__file__), 'files')
    for filename in os.listdir(files_dir):
        if filename.endswith('.docx'):
            filepath = os.path.join(files_dir, filename)
            print(f'Processing: {filename}')
            process_docx_file(filepath)
    print('处理完成！')

if __name__ == '__main__':
    main()
